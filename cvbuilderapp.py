from docx import Document
from docx.shared import Inches

document = Document()

# profile picture
document.add_picture('me.jpg', width=Inches(2.0))

# name phone number and email details
name = input('What is your name? ')
phone_number = input('What is your phone number? ')
email = input('What is your email? ')

document.add_paragraph(name.capitalize() + ' | ' + phone_number + ' | ' + email)

# about me
document.add_heading('About Me')
about_me = input('Tell me about yourself: ')
document.add_paragraph(about_me)

# experience one
document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter company: ')
from_date = input('Enter from date: ')
to_date = input('Enter to date')
position_roles = input('Describe your roles: ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + ' - ' + to_date + '\n').italic = True
p.add_run(position_roles)

# more experiences
while True:
    has_more_experiences = input('Do you have more experiences? \'Yes\' or \'No\'')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter company: ')
        from_date = input('Enter from date: ')
        to_date = input('Enter to date')
        position_roles = input('Describe your roles: ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + ' - ' + to_date + '\n').italic = True
        p.add_run(position_roles)

    else:
        break

# skills
document.add_heading('My Skills')
skill = input('Enter skill 1: ')
p_skills = document.add_paragraph(skill)
p_skills.style = 'List Bullet'

#more skills
while True:
    has_more_skills = input('Do you have more skills? \'Yes\' or \'No\'')
    if has_more_skills.lower() == 'yes':
        skill = input('Enter skill 1: ')
        p_skills = document.add_paragraph(skill)
        p_skills.style = 'List Bullet'
    else:
        break

#footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = 'Made in Kenya!'

document.save('cv.docx')
